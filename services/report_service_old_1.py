from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class ReportService:
    def __init__(self, output_dir="./data/reports"):

        self.output_dir = Path(output_dir)

        self.output_dir.mkdir(
            parents=True,
            exist_ok=True,
        )

    # =========================================================
    # PUBLIC METHOD
    # =========================================================

    def create_client_report(
        self,
        state,
        session_id,
    ):

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        filename = f"patent_analysis_{session_id}_{timestamp}.docx"

        report_path = self.output_dir / filename

        self._build_document(
            state,
            session_id,
            report_path,
        )

        return str(report_path)

    # =========================================================
    # BUILD DOCUMENT
    # =========================================================

    def _build_document(
        self,
        state,
        session_id,
        report_path,
    ):

        invention = state["invention"]
        keywords = state["keywords"]
        cpc_analysis = state["cpc_analysis"]

        document = Document()

        self._configure_page(document)

        self._add_header(document)

        self._add_footer(document)

        self._add_cover(
            document,
            invention,
            session_id,
        )

        document.add_page_break()

        # -----------------------------------------------------
        # 1. Executive Summary
        # -----------------------------------------------------

        self._add_executive_summary(
            document,
            invention,
        )

        # -----------------------------------------------------
        # 2. Invention Decomposition
        # -----------------------------------------------------

        self._add_invention_decomposition(
            document,
            invention,
        )

        # -----------------------------------------------------
        # 3. Search Vocabulary
        # -----------------------------------------------------

        self._add_search_vocabulary(
            document,
            keywords,
        )

        # -----------------------------------------------------
        # 4. CPC Analysis
        # -----------------------------------------------------

        self._add_cpc_analysis(
            document,
            cpc_analysis,
        )

        # -----------------------------------------------------
        # 5. Methodology
        # -----------------------------------------------------

        self._add_methodology(
            document,
        )

        # -----------------------------------------------------
        # 6. Traceability
        # -----------------------------------------------------

        self._add_traceability(
            document,
            session_id,
        )

        # -----------------------------------------------------
        # Confidentiality
        # -----------------------------------------------------

        self._add_confidentiality_notice(
            document,
        )

        document.save(report_path)

    # =========================================================
    # PAGE CONFIGURATION
    # =========================================================

    def _configure_page(
        self,
        document,
    ):

        section = document.sections[0]

        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        document.styles["Normal"].font.name = "Aptos"
        document.styles["Normal"].font.size = Pt(10)

    # =========================================================
    # HEADER
    # =========================================================

    def _add_header(
        self,
        document,
    ):

        header = document.sections[0].header

        paragraph = header.paragraphs[0]

        paragraph.text = "GVG Consulting Services | Patent Invention Analysis"

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for run in paragraph.runs:
            run.font.size = Pt(8)

    # =========================================================
    # FOOTER
    # =========================================================

    def _add_footer(
        self,
        document,
    ):

        footer = document.sections[0].footer

        paragraph = footer.paragraphs[0]

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("Confidential – Prepared for Client")

        run.font.size = Pt(8)

    # =========================================================
    # COVER
    # =========================================================

    def _add_cover(
        self,
        document,
        invention,
        session_id,
    ):

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("PATENT INVENTION ANALYSIS REPORT")

        run.bold = True
        run.font.size = Pt(24)

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("Invention Decomposition • Search Vocabulary • CPC Classification")

        run.italic = True
        run.font.size = Pt(12)

        document.add_paragraph()

        table = document.add_table(
            rows=4,
            cols=2,
        )

        table.style = "Table Grid"

        rows = [
            (
                "Client / Matter",
                "[CLIENT / MATTER NAME]",
            ),
            (
                "Session / Reference",
                session_id,
            ),
            (
                "Invention Title",
                getattr(
                    invention,
                    "title",
                    "[INVENTION TITLE]",
                ),
            ),
            (
                "CPC Corpus Version",
                "2026.08",
            ),
        ]

        for row, values in zip(
            table.rows,
            rows,
        ):
            self._set_cell(
                row.cells[0],
                values[0],
                bold=True,
            )

            self._shade_cell(
                row.cells[0],
                "E7E6E6",
            )

            self._set_cell(
                row.cells[1],
                values[1],
            )

        document.add_paragraph()

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("CONFIDENTIAL CLIENT DELIVERABLE")

        run.bold = True

    # =========================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================

    def _add_executive_summary(
        self,
        document,
        invention,
    ):

        document.add_heading(
            "1. Executive Summary",
            level=1,
        )

        document.add_paragraph(invention.invention_summary)

    # =========================================================
    # 2. INVENTION DECOMPOSITION
    # =========================================================

    def _add_invention_decomposition(
        self,
        document,
        invention,
    ):

        document.add_heading(
            "2. Invention Decomposition",
            level=1,
        )

        document.add_paragraph(
            "The invention has been decomposed into "
            "structural, procedural/process and functional "
            "features to support subsequent patent searching "
            "and classification."
        )

        # -----------------------------------------------------
        # Structural
        # -----------------------------------------------------

        self._add_feature_table(
            document,
            "2.1 Structural Features",
            invention.structural_features,
            "Structural",
        )

        # -----------------------------------------------------
        # Procedural
        # -----------------------------------------------------

        self._add_feature_table(
            document,
            "2.2 Procedural / Process Features",
            invention.procedural_features,
            "Procedural",
        )

        # -----------------------------------------------------
        # Functional
        # -----------------------------------------------------

        self._add_feature_table(
            document,
            "2.3 Functional Features",
            invention.functional_features,
            "Functional",
        )

    # =========================================================
    # FEATURE TABLE
    # =========================================================

    def _add_feature_table(
        self,
        document,
        heading,
        features,
        feature_type,
    ):

        document.add_heading(
            heading,
            level=2,
        )

        table = document.add_table(
            rows=1,
            cols=3,
        )

        table.style = "Table Grid"

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [
            "Feature ID",
            "Feature",
            "Technical Details",
        ]

        for index, header in enumerate(headers):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
            )

            self._shade_cell(
                table.rows[0].cells[index],
                "D9EAF7",
            )

        for item in features:
            row = table.add_row()

            self._set_cell(
                row.cells[0],
                item.feature_id,
            )

            # ---------------------------------------------
            # StructuralFeature
            # ---------------------------------------------

            if feature_type == "Structural":
                description = item.feature

            # ---------------------------------------------
            # ProceduralFeature
            # ---------------------------------------------

            elif feature_type == "Procedural":
                description = item.step

            # ---------------------------------------------
            # FunctionalFeature
            # ---------------------------------------------

            elif feature_type == "Functional":
                description = item.function

            else:
                description = ""

            self._set_cell(
                row.cells[1],
                description,
            )

            self._set_cell(
                row.cells[2],
                item.details,
            )

        document.add_paragraph()

    # =========================================================
    # 3. SEARCH VOCABULARY
    # =========================================================

    def _add_search_vocabulary(
        self,
        document,
        keywords,
    ):

        document.add_heading(
            "3. Search Vocabulary",
            level=1,
        )

        document.add_paragraph(
            "Technical search terminology was generated "
            "independently for structural, procedural/process "
            "and functional features and then consolidated."
        )

        # -----------------------------------------------------
        # Structural
        # -----------------------------------------------------

        self._add_keyword_table(
            document,
            "3.1 Structural Keywords",
            keywords.structural_keywords,
        )

        # -----------------------------------------------------
        # Procedural
        # -----------------------------------------------------

        self._add_keyword_table(
            document,
            "3.2 Procedural / Process Keywords",
            keywords.procedural_keywords,
        )

        # -----------------------------------------------------
        # Functional
        # -----------------------------------------------------

        self._add_keyword_table(
            document,
            "3.3 Functional Keywords",
            keywords.functional_keywords,
        )

        # -----------------------------------------------------
        # Consolidated
        # -----------------------------------------------------

        document.add_heading(
            "3.4 Consolidated Search Vocabulary",
            level=2,
        )

        table = document.add_table(
            rows=1,
            cols=2,
        )

        table.style = "Table Grid"

        self._set_cell(
            table.rows[0].cells[0],
            "No.",
            bold=True,
        )

        self._set_cell(
            table.rows[0].cells[1],
            "Consolidated Keyword / Term",
            bold=True,
        )

        self._shade_cell(
            table.rows[0].cells[0],
            "D9EAF7",
        )

        self._shade_cell(
            table.rows[0].cells[1],
            "D9EAF7",
        )

        for index, keyword in enumerate(
            keywords.consolidated_keywords,
            start=1,
        ):
            row = table.add_row()

            self._set_cell(
                row.cells[0],
                index,
            )

            self._set_cell(
                row.cells[1],
                keyword,
            )

        document.add_paragraph()

    # =========================================================
    # KEYWORD TABLE
    # =========================================================

    def _add_keyword_table(
        self,
        document,
        heading,
        keyword_sets,
    ):

        document.add_heading(
            heading,
            level=2,
        )

        table = document.add_table(
            rows=1,
            cols=4,
        )

        table.style = "Table Grid"

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [
            "Feature ID",
            "Feature",
            "Primary Keywords",
            "Synonyms / Alternative Terms",
        ]

        for index, header in enumerate(headers):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
            )

            self._shade_cell(
                table.rows[0].cells[index],
                "D9EAF7",
            )

        for item in keyword_sets:
            row = table.add_row()

            self._set_cell(
                row.cells[0],
                item.feature_id,
            )

            self._set_cell(
                row.cells[1],
                item.feature,
            )

            self._set_cell(
                row.cells[2],
                "; ".join(item.primary_keywords),
            )

            alternatives = item.synonyms + item.alternative_terms

            self._set_cell(
                row.cells[3],
                "; ".join(alternatives),
            )

        document.add_paragraph()

    # =========================================================
    # 4. CPC ANALYSIS
    # =========================================================

    def _add_cpc_analysis(
        self,
        document,
        cpc_analysis,
    ):

        document.add_heading(
            "4. CPC Classification Analysis",
            level=1,
        )

        document.add_paragraph(
            "Candidate CPC classifications were retrieved "
            "from the CPC 2026.08 corpus using semantic "
            "similarity and subsequently evaluated for "
            "technical relevance."
        )

        table = document.add_table(
            rows=1,
            cols=5,
        )

        table.style = "Table Grid"

        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [
            "CPC Code",
            "Classification / Title",
            "Feature(s)",
            "Relevance",
            "Technical Rationale",
        ]

        for index, header in enumerate(headers):
            self._set_cell(
                table.rows[0].cells[index],
                header,
                bold=True,
            )

            self._shade_cell(
                table.rows[0].cells[index],
                "D9EAD3",
            )

        for candidate in cpc_analysis.candidates:
            row = table.add_row()

            self._set_cell(
                row.cells[0],
                candidate.cpc_code,
            )

            self._set_cell(
                row.cells[1],
                candidate.title,
            )

            self._set_cell(
                row.cells[2],
                candidate.feature_id or "",
            )

            self._set_cell(
                row.cells[3],
                candidate.relevance,
            )

            self._set_cell(
                row.cells[4],
                candidate.reason,
            )

        document.add_paragraph()

    # =========================================================
    # 5. METHODOLOGY
    # =========================================================

    def _add_methodology(
        self,
        document,
    ):

        document.add_heading(
            "5. Methodology and Scope",
            level=1,
        )

        points = [
            (
                "The client-provided IDF was converted into "
                "a structured representation of structural, "
                "procedural/process and functional features."
            ),
            (
                "Technical keywords, synonyms and alternative "
                "terminology were generated independently "
                "for each feature category and then consolidated."
            ),
            ("CPC candidates were retrieved using semantic similarity against the CPC 2026.08 corpus."),
            ("Retrieved CPC candidates were evaluated for technical relevance using the language model."),
            (
                "This report represents an analytical "
                "search-preparation and classification stage "
                "and is not, by itself, a legal opinion on "
                "novelty, inventive step or patentability."
            ),
        ]

        for point in points:
            document.add_paragraph(
                point,
                style="List Bullet",
            )

    # =========================================================
    # 6. TRACEABILITY
    # =========================================================

    def _add_traceability(
        self,
        document,
        session_id,
    ):

        document.add_heading(
            "6. Traceability",
            level=1,
        )

        table = document.add_table(
            rows=5,
            cols=2,
        )

        table.style = "Table Grid"

        values = [
            (
                "Session ID",
                session_id,
            ),
            (
                "Analysis Date",
                datetime.now().isoformat(timespec="seconds"),
            ),
            (
                "CPC Version",
                "2026.08",
            ),
            (
                "Analysis Stages",
                "Decomposition → Keywords → CPC Retrieval → CPC Evaluation",
            ),
            (
                "Report Status",
                "Generated from structured graph state",
            ),
        ]

        for row, values in zip(
            table.rows,
            values,
        ):
            self._set_cell(
                row.cells[0],
                values[0],
                bold=True,
            )

            self._shade_cell(
                row.cells[0],
                "E7E6E6",
            )

            self._set_cell(
                row.cells[1],
                values[1],
            )

    # =========================================================
    # CONFIDENTIALITY
    # =========================================================

    def _add_confidentiality_notice(
        self,
        document,
    ):

        document.add_paragraph()

        paragraph = document.add_paragraph()

        run = paragraph.add_run("Confidentiality Notice")

        run.bold = True

        document.add_paragraph(
            "This report is intended for the client and "
            "authorized representatives only. It should "
            "be reviewed by the patent practitioner before "
            "being relied upon for any filing, prosecution, "
            "freedom-to-operate, validity or other legal "
            "decision."
        )

    # =========================================================
    # TABLE HELPERS
    # =========================================================

    @staticmethod
    def _set_cell(
        cell,
        text,
        bold=False,
    ):

        cell.text = ""

        paragraph = cell.paragraphs[0]

        run = paragraph.add_run(str(text))

        run.bold = bold
        run.font.size = Pt(9)

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    @staticmethod
    def _shade_cell(
        cell,
        fill,
    ):

        tcPr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement("w:shd")

        shd.set(
            qn("w:fill"),
            fill,
        )

        tcPr.append(shd)
